#!/usr/bin/env python3
"""Convert ota-design.md → ota-design.docx"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── styles ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

# Heading styles
for i, sz, clr in [(1, 18, '1d1d1f'), (2, 14, '0071e3'), (3, 12, '333333'), (4, 11, '555555')]:
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.font.size = Pt(sz)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor.from_string(clr)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.paragraph_format.space_before = Pt(12 if i < 3 else 8)
    hs.paragraph_format.space_after = Pt(6)

with open('docs/ota-design.md', 'rb') as f:
    raw = f.read()
# Strip control characters except \n \r \t
raw = re.sub(rb'[\x00-\x08\x0b\x0c\x0e-\x1f]', b'', raw)
text = raw.decode('utf-8', errors='replace')
lines = text.split('\n')

i = 0
in_code = False
in_table = False
code_lines = []
table_rows = []

while i < len(lines):
    line = lines[i].rstrip()

    # ── code block ──
    if line.startswith('```'):
        if in_code:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string('333333')
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            code_lines = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # ── table ──
    if '|---' in line and i+1 < len(lines) and '|---' in lines[i-1].rstrip():
        # Previous line was header
        i += 1
        continue

    table_match = re.match(r'^\|(.+)\|$', line)
    next_is_table = i+1 < len(lines) and re.match(r'^\|(.+)\|$', lines[i+1].rstrip())
    prev_was_table = i > 0 and re.match(r'^\|(.+)\|$', lines[i-1].rstrip())

    if table_match and (next_is_table or prev_was_table):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        table_rows.append(cells)

        # If next line is NOT a table, flush
        if not next_is_table and table_rows:
            # Skip separator row (---)
            data_rows = []
            for r in table_rows:
                if any(re.match(r'^[-:]+$', c) for c in r):
                    continue
                data_rows.append(r)

            if len(data_rows) >= 2:
                tbl = doc.add_table(rows=len(data_rows), cols=len(data_rows[0]))
                tbl.style = 'Light Shading Accent 1'
                for ri, row in enumerate(data_rows):
                    for ci, cell in enumerate(row):
                        tbl.rows[ri].cells[ci].text = cell

            table_rows = []
        i += 1
        continue

    # ── horizontal rule ──
    if line.strip().startswith('---') and len(line.strip()) <= 5:
        doc.add_paragraph('─' * 60)
        i += 1
        continue

    # ── heading ──
    h1 = re.match(r'^#\s+(.+)', line)
    h2 = re.match(r'^##\s+(.+)', line)
    h3 = re.match(r'^###\s+(.+)', line)
    h4 = re.match(r'^####\s+(.+)', line)

    if h1:
        doc.add_heading(h1.group(1), level=1)
    elif h2:
        doc.add_heading(h2.group(1), level=2)
    elif h3:
        doc.add_heading(h3.group(1), level=3)
    elif h4:
        doc.add_heading(h4.group(1), level=4)
    else:
        # ── regular paragraph with inline formatting ──
        text = line.strip()
        if not text:
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            # Bold
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # > quotes
                    part = re.sub(r'^>\s?', '', part)
                    run = p.add_run(part)

    i += 1

doc.save('docs/ota-design.docx')
print('done → docs/ota-design.docx')
